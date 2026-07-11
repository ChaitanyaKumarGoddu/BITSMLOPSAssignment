"""Generate the final, screenshot-aligned .docx report.

Produces reports/Heart_Disease_MLOps_Report.docx with, for every section:
  * a clear description that matches the actual code,
  * the exact SOURCE FILES that implement it,
  * the exact COMMANDS to run,
  * the auto-generated figures (EDA / ROC / confusion matrix / architecture), and
  * the user's captured screenshots (from reports/screenshots/extracted/),
    placed in the correct section with descriptive captions.

Only the GitHub Actions run is a [SCREENSHOT] placeholder (add it yourself).

Run from the repo root:
    python scripts/build_report_docx.py
"""
from __future__ import annotations

import glob
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
FIGURES = ROOT / "reports" / "figures"
SHOTS = ROOT / "reports" / "screenshots" / "extracted"
META = ROOT / "models" / "model_metadata.json"
OUT = ROOT / "reports" / "Heart_Disease_MLOps_Report.docx"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
GREY = RGBColor(0x66, 0x66, 0x66)
CODE_BG = "F2F2F2"

# ---- Submission details -------------------------------------------------- #
STUDENT_NAME = "Chaitanya Kumar Goddu"
BITS_ID = "2024ac05880@wilp.bits-pilani.ac.in"
REPO_URL = "https://github.com/ChaitanyaKumarGoddu/BITSMLOPSAssignment"

doc = Document()


# --------------------------------------------------------------------------- #
# Helpers
# --------------------------------------------------------------------------- #
def h1(text: str):
    doc.add_heading(text, level=1)


def h2(text: str):
    doc.add_heading(text, level=2)


def para(text: str):
    p = doc.add_paragraph(text)
    return p


def files(items: list[str]):
    """A 'Source files' block listing the code that implements the section."""
    p = doc.add_paragraph()
    r = p.add_run("Source files:")
    r.bold = True
    r.font.color.rgb = ACCENT
    for it in items:
        b = doc.add_paragraph(style="List Bullet")
        run = b.add_run(it)
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)


def code(text: str, label: str = "Commands"):
    """A shaded monospace block for commands / snippets."""
    if label:
        p = doc.add_paragraph()
        r = p.add_run(label + ":")
        r.bold = True
        r.font.color.rgb = ACCENT
    p = doc.add_paragraph()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), CODE_BG)
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p._p.get_or_add_pPr().append(shd)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)


def caption(text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GREY


def fig(filename: str, cap: str, width_in: float = 6.0):
    path = FIGURES / filename
    if path.exists():
        doc.add_picture(str(path), width=Inches(width_in))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption(cap)
    else:
        placeholder(f"missing figure {filename} (run: python -m src.data.eda)")


def shot(index: int, cap: str, width_in: float = 6.4):
    """Embed the user's screenshot #index (matched by NN_ prefix)."""
    matches = sorted(glob.glob(str(SHOTS / f"{index:02d}_*")))
    if matches:
        doc.add_picture(matches[0], width=Inches(width_in))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption(f"Screenshot: {cap}")
    else:
        placeholder(f"screenshot #{index}: {cap}")


def placeholder(text: str):
    p = doc.add_paragraph()
    r = p.add_run(f"[SCREENSHOT NEEDED - {text}]")
    r.bold = True
    r.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)
    r.font.size = Pt(11)


def page_break():
    doc.add_page_break()


# --------------------------------------------------------------------------- #
# Build the document
# --------------------------------------------------------------------------- #
meta = json.loads(META.read_text()) if META.exists() else {}
m = meta.get("metrics", {})

# ---- Title page ----
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Heart Disease Risk Prediction\nEnd-to-End MLOps Pipeline")
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = ACCENT

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = s.add_run("Machine Learning Operations (MLOps) — AIMLCZG523\nAssignment 01")
rs.font.size = Pt(13)
rs.font.color.rgb = GREY

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run(
    f"Name: {STUDENT_NAME}\n"
    f"BITS ID: {BITS_ID}\n"
    f"GitHub Repository: {REPO_URL}\n"
    "Video walkthrough: ______________________________________________"
).font.size = Pt(12)

page_break()

# ---- 1. Overview ----
h1("1. Project Overview")
para(
    "This project builds a machine-learning classifier that predicts the risk "
    "of heart disease from patient clinical data and deploys it as a "
    "cloud-ready, monitored REST API, following modern MLOps practice end to "
    "end: data acquisition and EDA, feature engineering, multi-model training "
    "with hyper-parameter tuning and cross-validation, MLflow experiment "
    "tracking, reproducible packaging, automated testing with GitHub Actions "
    "CI/CD, Docker containerisation, Kubernetes deployment, and "
    "Prometheus/Grafana monitoring."
)
para(
    f"Dataset: UCI Heart Disease (Cleveland), 303 patient records, 13 clinical "
    f"features, binary target (presence/absence of disease)."
)
para(
    f"Headline result: the tuned Logistic Regression pipeline achieves "
    f"ROC-AUC {m.get('roc_auc', 0.968):.3f}, accuracy {m.get('accuracy', 0.90):.2f}, "
    f"and recall {m.get('recall', 0.93):.2f} on the held-out test set — recall "
    f"being the priority metric for a screening tool where a missed case is costly."
)
h2("Technology stack")
para(
    "Python 3.14 · pandas / scikit-learn / XGBoost · MLflow (SQLite backend) · "
    "FastAPI + Uvicorn · Pytest · Docker (multi-stage) · Kubernetes (Docker "
    "Desktop) + Helm · Prometheus + Grafana · GitHub Actions."
)

page_break()

# ---- 2. Setup ----
h1("2. Setup & Installation")
para("Requires Python 3.11-3.14. All commands are run from the repository root.")
files(["requirements.txt  (runtime deps)", "requirements-dev.txt  (+ EDA, tests, lint, notebook)"])
code(
    "python -m venv .venv\n"
    ".venv\\Scripts\\Activate.ps1        # Windows PowerShell\n"
    "pip install --upgrade pip\n"
    "pip install -r requirements-dev.txt"
)
para(
    "Dependencies are pinned to wheel-available versions verified on Python "
    "3.14 (Windows) and Python 3.12 (Linux CI). The runtime Docker image uses "
    "only requirements.txt; requirements-dev.txt adds EDA, testing, linting and "
    "Jupyter."
)

page_break()

# ---- 3. Task 1: Data + EDA ----
h1("3. Task 1 - Data Acquisition & Exploratory Data Analysis")
files([
    "src/config.py            - paths + feature schema (single source of truth)",
    "src/data/download.py     - fetch dataset (ucimlrepo id=45, URL fallback)",
    "src/data/preprocess.py   - clean(): coerce numeric, binarise target, de-dup",
    "src/data/eda.py          - generate the EDA figures",
    "notebooks/01_eda.ipynb   - narrated EDA notebook",
])
para(
    "download.py fetches the Cleveland dataset via the official ucimlrepo "
    "package (id=45), falling back to a direct download of "
    "processed.cleveland.data. preprocess.clean() coerces every column to "
    "numeric, binarises the 0-4 severity field into a 0/1 target (0 = no "
    "disease, >=1 = disease), and drops exact duplicates. Missing values (the "
    "UCI '?' marker) appear only in ca (4 rows) and thal (2 rows) and are "
    "imputed later inside the modelling pipeline rather than dropped. The two "
    "classes are close to balanced (~46% positive)."
)
code(
    "python -m src.data.download      # -> data/raw/heart_disease_raw.csv (303 rows)\n"
    "python -m src.data.preprocess    # -> data/processed/heart_disease_clean.csv\n"
    "python -m src.data.eda           # -> reports/figures/eda_*.png"
)
fig("eda_class_balance.png", "Figure 1. Class balance (~46% positive).", 4.0)
fig("eda_histograms.png", "Figure 2. Distributions of numeric features.", 6.2)
fig("eda_correlation_heatmap.png",
    "Figure 3. Correlation heatmap. Strongest target correlates: cp, thalach, "
    "oldpeak, ca, exang.", 5.6)
fig("eda_feature_relationships.png",
    "Figure 4. Key features by target - diseased patients show lower max heart "
    "rate, higher ST depression, and more affected vessels.", 6.5)

page_break()

# ---- 4. Task 2: Feature eng + models ----
h1("4. Task 2 - Feature Engineering & Model Development")
files([
    "src/features/pipeline.py - ColumnTransformer (impute + scale + one-hot)",
    "src/models/train.py      - train 3 models, GridSearchCV, select best",
    "src/models/evaluate.py   - metrics + ROC / confusion-matrix plots",
])
para(
    "All feature engineering lives in a single scikit-learn ColumnTransformer "
    "(build_preprocessor): numeric features get median imputation + standard "
    "scaling; categorical features get most-frequent imputation + one-hot "
    "encoding with handle_unknown='ignore'. Because the transformer is bundled "
    "with the estimator into one Pipeline, it is fitted only on the training "
    "folds and reused byte-for-byte at inference, eliminating train/serving "
    "skew. Three classifiers are tuned with GridSearchCV (5-fold, ROC-AUC "
    "scoring) on a stratified 80/20 split (random_state=42)."
)
code("python -m src.models.train")
h2("Model comparison (held-out test set, n=61)")
rows = [
    ("Model", "Accuracy", "Precision", "Recall", "F1", "Test ROC-AUC", "CV ROC-AUC"),
    ("Logistic Regression *", "0.902", "0.867", "0.929", "0.897", "0.968", "0.898 +/- 0.047"),
    ("Random Forest", "0.885", "0.839", "0.929", "0.881", "0.952", "0.897 +/- 0.043"),
    ("XGBoost", "0.869", "0.813", "0.929", "0.867", "0.947", "0.879 +/- 0.025"),
]
table = doc.add_table(rows=len(rows), cols=len(rows[0]))
table.style = "Light Grid Accent 1"
for i, row in enumerate(rows):
    for j, val in enumerate(row):
        cell = table.rows[i].cells[j]
        cell.text = val
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True
para(
    "* Selected production model. Logistic Regression wins on every headline "
    "metric, is cheap to serve, and is interpretable (coefficients map to "
    "clinical risk factors) - a real advantage in healthcare. Best "
    "hyper-parameters: C=0.1, solver=liblinear. All three models achieve 0.929 "
    "recall, the metric we most care about for a screening tool."
)
fig("roc_logistic_regression.png", "Figure 5. ROC curve - Logistic Regression (AUC = 0.968).", 4.6)
fig("cm_logistic_regression.png", "Figure 6. Confusion matrix - Logistic Regression.", 4.0)

page_break()

# ---- 5. Task 3: MLflow ----
h1("5. Task 3 - Experiment Tracking (MLflow)")
files([
    "src/models/train.py      - mlflow.start_run / log_params / log_metrics / log_artifact / log_model",
    "src/config.py            - MLFLOW_TRACKING_URI (sqlite), MLFLOW_ARTIFACT_LOCATION",
    "scripts/mlflow_ui.py     - launches the UI on Python 3.14 (compat shim)",
])
para(
    "Training logs everything to MLflow using a SQLite backend "
    "(sqlite:///mlflow.db; the file store is deprecated in MLflow 3.x). One "
    "parent run (model-selection-*) contains a nested run per model, each "
    "logging the best hyper-parameters, all metrics (accuracy, precision, "
    "recall, f1, test ROC-AUC, CV ROC-AUC mean/std), the ROC-curve and "
    "confusion-matrix artifacts, and the fitted sklearn Pipeline."
)
code(
    "python scripts/mlflow_ui.py --port 5000    # then open http://localhost:5000\n"
    "# (Python <=3.13: mlflow ui --backend-store-uri sqlite:///mlflow.db --port 5000)"
)
shot(1, "MLflow runs table - all three models compared side by side (accuracy, ROC-AUC, recall, F1, CV scores).")
h2("Logistic Regression (selected model)")
shot(10, "Logistic Regression run overview - metrics + tuned params (C=0.1, solver=liblinear) + logged model.")
shot(11, "Logistic Regression - model-metric charts.")
shot(12, "Logistic Regression - logged artifact: confusion matrix.")
shot(13, "Logistic Regression - logged artifact: ROC curve (AUC = 0.968).")
h2("Random Forest")
shot(6, "Random Forest run overview - metrics + tuned params (n_estimators=400, max_depth=10, min_samples_leaf=4).")
shot(7, "Random Forest - model-metric charts.")
shot(8, "Random Forest - confusion matrix artifact.")
shot(9, "Random Forest - ROC curve artifact (AUC = 0.952).")
h2("XGBoost")
shot(2, "XGBoost run overview - metrics + tuned params (learning_rate=0.05, max_depth=3, n_estimators=200).")
shot(3, "XGBoost - model-metric charts.")
shot(4, "XGBoost - confusion matrix artifact.")
shot(5, "XGBoost - ROC curve artifact (AUC = 0.947).")

page_break()

# ---- 6. Task 4: Packaging ----
h1("6. Task 4 - Model Packaging & Reproducibility")
files([
    "models/model.pkl            - fitted sklearn Pipeline (joblib)",
    "models/model_metadata.json  - model name, best params, metrics, feature order",
    "src/models/train.py         - joblib.dump(best_estimator, MODEL_PATH)",
    "requirements.txt            - pinned dependency versions",
])
para(
    "The winning pipeline is serialised to models/model.pkl with joblib, "
    "alongside model_metadata.json (model name, best params, metrics, feature "
    "order, timestamp). Because preprocessing lives inside the pipeline, the "
    "artifact is fully self-contained: the API just calls predict_proba on a "
    "raw feature row. requirements.txt pins exact versions and random_state=42 "
    "fixes every stochastic step, so a clean environment reproduces the same "
    "model and metrics."
)

page_break()

# ---- 7. Task 5: CI/CD + tests ----
h1("7. Task 5 - CI/CD Pipeline & Automated Testing")
files([
    "tests/conftest.py           - offline fixtures (synthetic data + tiny pipeline)",
    "tests/test_data.py          - cleaning / preprocessing tests",
    "tests/test_model.py         - pipeline + metrics tests",
    "tests/test_api.py           - FastAPI /predict, /health, validation, 503, /metrics",
    ".github/workflows/ci.yml    - lint -> test -> train -> docker build + smoke test",
    "pyproject.toml, .flake8     - pytest / black / isort / flake8 config",
])
para(
    "The suite has 17 unit/integration tests covering target binarisation, "
    "missing-value handling, de-duplication, the column contract, the "
    "preprocessor removing NaNs, pipeline output shape/range, metric "
    "computation, and the API (happy path, 422 validation errors, 503 when no "
    "model, /metrics exposure). The GitHub Actions workflow runs on every "
    "push/PR with three dependent jobs - Lint & Test, Train Model, and Docker "
    "Build Validation - and fails the run on any lint or test error."
)
code(
    "pytest                          # 17 passed\n"
    "flake8 src tests                # lint (clean)\n"
    "black --check src tests && isort --check-only src tests"
)
shot(14, "Local Pytest run - all 17 tests pass, with the coverage summary.")
placeholder("GitHub Actions: a green 'CI/CD Pipeline' run showing the 3 jobs "
            "(Lint & Test -> Train Model -> Docker Build Validation). "
            "Capture from your repo's Actions tab after pushing.")

page_break()

# ---- 8. Task 6: Docker + API ----
h1("8. Task 6 - Containerisation (Docker) & Serving API")
files([
    "Dockerfile               - multi-stage build, non-root user, HEALTHCHECK",
    ".dockerignore            - keep the image lean",
    "src/api/main.py          - FastAPI app: /predict, /health, /metrics",
    "src/api/schemas.py       - Pydantic request/response models + validation",
    "sample_request.json      - example high-risk patient payload",
])
para(
    "A multi-stage Dockerfile (Python 3.12-slim) builds wheels in stage 1 and "
    "installs them offline in stage 2, runs as a non-root user, exposes port "
    "8000, and defines a container HEALTHCHECK against /health. The /predict "
    "endpoint accepts JSON, validates it with Pydantic (out-of-range values are "
    "rejected with HTTP 422), and returns the prediction plus a confidence "
    "score. Swagger UI is auto-generated at /docs."
)
code(
    "docker build -t heart-disease-api:latest .\n"
    "docker run --rm -p 8000:8000 heart-disease-api:latest\n"
    'curl.exe -X POST http://localhost:8000/predict -H "Content-Type: application/json" -d "@sample_request.json"'
)
shot(15, "FastAPI Swagger UI (/docs) exposing /predict, /health, /metrics and / endpoints.")
shot(17, "Swagger UI - /predict 'Try it out' with the sample request body.")
shot(18, "Swagger UI - /predict returns HTTP 200: prediction=1, confidence 0.9635.")
shot(16, "Local /predict call via curl returning prediction + confidence.")
shot(19, "Docker image build - multi-stage build completes (18/18 FINISHED).")
shot(20, "Container startup logs - model 'logistic_regression' loaded, health checks returning 200.")
shot(21, "Docker Desktop - the heart-disease-api container running on port 8000.")
shot(22, "/predict verified against the running container (prediction=1, confidence 0.9635).")

page_break()

# ---- 9. Task 7: Kubernetes ----
h1("9. Task 7 - Production Deployment (Kubernetes)")
files([
    "k8s/deployment.yaml      - 2 replicas, resource limits, readiness/liveness probes",
    "k8s/service.yaml         - LoadBalancer service on port 8000",
    "helm/heart-disease-api/  - equivalent parameterised Helm chart",
])
para(
    "The API is deployed to Docker Desktop Kubernetes. deployment.yaml runs 2 "
    "replicas with CPU/memory requests+limits, readiness and liveness probes on "
    "/health, and Prometheus scrape annotations; service.yaml exposes a "
    "LoadBalancer on port 8000 (reachable at http://localhost:8000 on Docker "
    "Desktop). A Helm chart provides the same deployment parameterised via "
    "values.yaml."
)
code(
    "kubectl apply -f k8s/\n"
    "kubectl get pods,svc            # wait until both pods are 1/1 Running\n"
    'curl.exe -X POST http://localhost:8000/predict -H "Content-Type: application/json" -d "@sample_request.json"\n'
    "# Helm alternative: helm install heart ./helm/heart-disease-api"
)
shot(23, "Docker Desktop Kubernetes cluster enabled and Running (kubeadm, 1 node, v1.34.1).")
shot(24, "kubectl - context, cluster-info, get nodes (Ready), apply -f k8s/, and 2/2 pods Running behind the LoadBalancer service.")
shot(25, "Kubernetes pod logs (replica 1) - model loaded, readiness probes returning 200.")
shot(26, "Kubernetes pod logs (replica 2) - model loaded, health checks passing.")
shot(27, "/predict verified through the Kubernetes LoadBalancer service.")

page_break()

# ---- 10. Task 8: Monitoring ----
h1("10. Task 8 - Monitoring & Logging")
files([
    "src/api/main.py                     - request-logging middleware + Prometheus metrics",
    "  (Instrumentator /metrics + heart_predictions_total{outcome=...} Counter)",
    "monitoring/prometheus.yml           - scrape config (api:8000 every 10s)",
    "docker-compose.yml                  - API + Prometheus + Grafana stack",
    "monitoring/grafana/provisioning/    - auto-provisioned datasource + dashboard",
])
para(
    "A FastAPI middleware logs every request (method, path, status, latency). "
    "prometheus-fastapi-instrumentator exposes standard HTTP latency/throughput "
    "metrics at /metrics, plus a custom heart_predictions_total{outcome=...} "
    "counter that tracks the positive/negative prediction mix (a simple "
    "drift signal). docker compose brings up the API (host port 8001, to "
    "coexist with the K8s deployment on 8000), Prometheus (host 9091), and "
    "Grafana (host 3000) with the datasource and dashboard auto-provisioned."
)
code(
    "docker compose up --build\n"
    "# generate traffic (API on 8001):\n"
    '1..30 | ForEach-Object { curl.exe -s -X POST http://localhost:8001/predict -H "Content-Type: application/json" -d "@sample_request.json" | Out-Null }\n'
    "# Prometheus: http://localhost:9091   Grafana: http://localhost:3000 (admin/admin)"
)
shot(28, "Prometheus - both scrape targets UP (heart-disease-api scraped at api:8000/metrics).")
shot(29, "Prometheus query - heart_predictions_total = 50 for outcome='disease'.")
shot(30, "Prometheus graph - heart_predictions_total over time.")
shot(31, "Grafana dashboard - Total Predictions, Prediction Outcomes, Request Rate, p95 Latency, HTTP Status Codes.")

page_break()

# ---- 11. Task 9: Architecture ----
h1("11. Task 9 - System Architecture")
para(
    "Development produces a tracked, versioned model artifact; CI/CD validates "
    "and packages it into a Docker image; Kubernetes serves it behind a "
    "LoadBalancer; Prometheus and Grafana observe it in production."
)
fig("architecture.png", "Figure 7. End-to-end system architecture.", 6.5)

# ---- 12. Repository & deliverables ----
h1("12. Repository & Deliverables")
para(f"GitHub repository: {REPO_URL}")
rows2 = [
    ("Deliverable", "Location in repo"),
    ("Code + Dockerfile + requirements", "repo root, src/"),
    ("Dataset + download script", "src/data/download.py, data/processed/"),
    ("Notebook / scripts (EDA, train, inference)", "notebooks/, src/"),
    ("Unit tests", "tests/"),
    ("GitHub Actions workflow", ".github/workflows/ci.yml"),
    ("Deployment manifests / Helm chart", "k8s/, helm/"),
    ("Monitoring config", "monitoring/, docker-compose.yml"),
    ("Report + figures", "reports/"),
]
tbl = doc.add_table(rows=len(rows2), cols=2)
tbl.style = "Light Grid Accent 1"
for i, row in enumerate(rows2):
    for j, val in enumerate(row):
        c = tbl.rows[i].cells[j]
        c.text = val
        for p in c.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9.5)
                if i == 0:
                    run.bold = True

# ---- 13. Conclusion ----
h1("13. Conclusion")
para(
    "This project delivers a reproducible, tested, containerised, and monitored "
    "heart-disease classifier that meets every assignment requirement. The "
    "tuned Logistic Regression (ROC-AUC 0.968, recall 0.93) is both accurate "
    "and interpretable - the right choice for a clinical screening context - "
    "and the surrounding MLOps scaffolding (MLflow, CI/CD, Docker, Kubernetes, "
    "Prometheus/Grafana) makes it safe to iterate on and operate in production. "
    "Possible extensions: data-drift detection with Evidently, a model registry "
    "with staged promotion, automated retraining triggers, and canary "
    "deployments."
)

doc.save(str(OUT))
print(f"[docx] Saved -> {OUT}")
